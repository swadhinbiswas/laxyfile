"""
Document Preview Renderer

This module provides preview capabilities for documents (PDF, DOC, etc.)
and archive files with content listing and metadata extraction.
"""

import time
import zipfile
import tarfile
import mimetypes
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
import tempfile
import subprocess

from .preview_system import BasePreviewRenderer, PreviewResult, PreviewType, PreviewConfig
from ..utils.logger import Logger


class DocumentPreviewRenderer(BasePreviewRenderer):
    """Renderer for document and archive files"""

    def __init__(self, config: PreviewConfig):
        super().__init__(config)

        # Check for optional dependencies
        self.pypdf_available = False
        self.docx_available = False
        self.openpyxl_available = False

        try:
            import PyPDF2
            self.PyPDF2 = PyPDF2
            self.pypdf_available = True
        except ImportError:
            try:
                import pypdf
                self.pypdf = pypdf
                self.pypdf_available = True
            except ImportError:
                self.logger.warning("PyPDF2/pypdf not available, PDF preview limited")

        try:
            from docx import Document
            self.docx_Document = Document
            self.docx_available = True
        except ImportError:
            self.logger.warning("python-docx not available, DOCX preview limited")

        try:
            import openpyxl
            self.openpyxl = openpyxl
            self.openpyxl_available = True
        except ImportError:
            self.logger.warning("openpyxl not available, Excel preview limited")

        # Supported formats
        self.document_extensions = {
            '.pdf', '.doc', '.docx', '.odt', '.rtf', '.txt', '.md', '.rst',
            '.xls', '.xlsx', '.ods', '.csv', '.ppt', '.pptx', '.odp'
        }

        self.archive_extensions = {
            '.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz', '.tgz', '.tbz2'
        }

    def can_prew(self, file_path: Path) -> bool:
        """Check if file is a document or archive"""
        extension = file_path.suffix.lower()
        return (extension in self.document_extensions or
                extension in self.archive_extensions)

    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return list(self.document_extensions | self.archive_extensions)

    def get_preview_type(self, file_path: Path) -> PreviewType:
        """Get preview type for file"""
        extension = file_path.suffix.lower()
        if extension in self.document_extensions:
            return PreviewType.DOCUMENT
        elif extension in self.archive_extensions:
            return PreviewType.ARCHIVE
        return PreviewType.UNKNOWN

    async def generate_preview(self, file_path: Path) -> PreviewResult:
        """Generate document/archive preview"""
        start_time = time.time()

        try:
            if self._is_file_too_large(file_path):
                return self._create_error_result(file_path, "File too large for preview")

            extension = file_path.suffix.lower()

            if extension in self.document_extensions:
                return await self._generate_document_preview(file_path, start_time)
            elif extension in self.archive_extensions:
                return await self._generate_archive_preview(file_path, start_time)
            else:
                return self._create_error_result(file_path, "Unsupported document format")

        except Exception as e:
            self.logger.error(f"Error generating document preview for {file_path}: {e}")
            return self._create_error_result(file_path, str(e))

    async def _generate_document_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate document preview"""
        try:
            extension = file_path.suffix.lower()

            if extension == '.pdf':
                return await self._generate_pdf_preview(file_path, start_time)
            elif extension in ['.docx']:
                return await self._generate_docx_preview(file_path, start_time)
            elif extension in ['.xlsx', '.xls']:
                return await self._generate_excel_preview(file_path, start_time)
            elif extension in ['.txt', '.md', '.rst']:
                return await self._generate_text_document_preview(file_path, start_time)
            elif extension == '.csv':
                return await self._generate_csv_preview(file_path, start_time)
            else:
                return await self._generate_generic_document_preview(file_path, start_time)

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_pdf_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate PDF preview"""
        try:
            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': 'application/pdf'
            }

            content = f"📄 PDF Document: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"

            if self.pypdf_available:
                try:
                    if hasattr(self, 'PyPDF2'):
                        # PyPDF2 version
                        with open(file_path, 'rb') as file:
                            pdf_reader = self.PyPDF2.PdfReader(file)

                            metadata.update({
                                'page_count': len(pdf_reader.pages),
                                'encrypted': pdf_reader.is_encrypted
                            })

                            # Extract document info
                            if pdf_reader.metadata:
                                doc_info = {}
                                for key, value in pdf_reader.metadata.items():
                                    if key.startswith('/'):
                                        key = key[1:]  # Remove leading slash
                                    doc_info[key] = str(value)
                                metadata['document_info'] = doc_info

                            content += f"Pages: {metadata['page_count']}\n"

                            if metadata['encrypted']:
                                content += "🔒 Document is encrypted\n"

                            # Extract text from first few pages
                            text_content = []
                            max_pages = min(self.config.pdf_max_pages, len(pdf_reader.pages))

                            for i in range(max_pages):
                                try:
                                    page = pdf_reader.pages[i]
                                    page_text = page.extract_text()
                                    if page_text.strip():
                                        text_content.append(f"\n--- Page {i+1} ---\n{page_text.strip()}")
                                except Exception as e:
                                    text_content.append(f"\n--- Page {i+1} ---\n[Error extracting text: {e}]")

                            if text_content:
                                content += "\n📝 Content Preview:\n" + "\n".join(text_content)

                                if len(pdf_reader.pages) > max_pages:
                                    content += f"\n\n... (showing first {max_pages} of {len(pdf_reader.pages)} pages)"

                    else:
                        # pypdf version
                        with open(file_path, 'rb') as file:
                            pdf_reader = self.pypdf.PdfReader(file)

                            metadata.update({
                                'page_count': len(pdf_reader.pages),
                                'encrypted': pdf_reader.is_encrypted
                            })

                            content += f"Pages: {metadata['page_count']}\n"

                            if metadata['encrypted']:
                                content += "🔒 Document is encrypted\n"

                except Exception as e:
                    content += f"\nError reading PDF: {e}"
            else:
                content += "\n(PDF library not available for detailed preview)"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_docx_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate DOCX preview"""
        try:
            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            }

            content = f"📝 Word Document: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"

            if self.docx_available:
                try:
                    doc = self.docx_Document(file_path)

                    # Extract document properties
                    core_props = doc.core_properties
                    if core_props:
                        doc_info = {}
                        if core_props.title:
                            doc_info['title'] = core_props.title
                        if core_props.author:
                            doc_info['author'] = core_props.author
                        if core_props.subject:
                            doc_info['subject'] = core_props.subject
                        if core_props.created:
                            doc_info['created'] = str(core_props.created)
                        if core_props.modified:
                            doc_info['modified'] = str(core_props.modified)

                        if doc_info:
                            metadata['document_info'] = doc_info

                    # Extract text content
                    paragraphs = []
                    char_count = 0

                    for paragraph in doc.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            paragraphs.append(text)
                            char_count += len(text)

                            # Limit content length
                            if char_count > self.config.max_text_length:
                                break

                    metadata.update({
                        'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
                        'character_count': char_count
                    })

                    content += f"Paragraphs: {metadata['paragraph_count']}\n"
                    content += f"Characters: {metadata['character_count']:,}\n"

                    if paragraphs:
                        content += "\n📝 Content Preview:\n"
                        preview_text = '\n\n'.join(paragraphs[:20])  # First 20 paragraphs

                        if len(preview_text) > self.config.max_text_length:
                            preview_text = preview_text[:self.config.max_text_length] + "..."

                        content += preview_text

                        if len(paragraphs) > 20:
                            content += f"\n\n... (showing first 20 paragraphs)"

                except Exception as e:
                    content += f"\nError reading DOCX: {e}"
            else:
                content += "\n(python-docx not available for detailed preview)"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_excel_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate Excel preview"""
        try:
            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            }

            content = f"📊 Excel Spreadsheet: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"

            if self.openpyxl_available:
                try:
                    workbook = self.openpyxl.load_workbook(file_path, read_only=True)

                    sheet_names = workbook.sheetnames
                    metadata.update({
                        'sheet_count': len(sheet_names),
                        'sheet_names': sheet_names
                    })

                    content += f"Sheets: {len(sheet_names)}\n"
                    content += f"Sheet Names: {', '.join(sheet_names)}\n"

                    # Preview first sheet
                    if sheet_names:
                        first_sheet = workbook[sheet_names[0]]

                        # Get dimensions
                        max_row = first_sheet.max_row
                        max_col = first_sheet.max_column

                        metadata.update({
                            'first_sheet_rows': max_row,
                            'first_sheet_columns': max_col
                        })

                        content += f"\nFirst Sheet '{sheet_names[0]}':\n"
                        content += f"Dimensions: {max_row} rows × {max_col} columns\n"

                        # Show first few rows
                        content += "\n📋 Data Preview:\n"
                        preview_rows = min(10, max_row)
                        preview_cols = min(5, max_col)

                        for row in range(1, preview_rows + 1):
                            row_data = []
                            for col in range(1, preview_cols + 1):
                                cell_value = first_sheet.cell(row=row, column=col).value
                                if cell_value is not None:
                                    row_data.append(str(cell_value)[:20])  # Truncate long values
                                else:
                                    row_data.append("")

                            content += f"Row {row}: {' | '.join(row_data)}\n"

                        if max_row > preview_rows or max_col > preview_cols:
                            content += f"... (showing first {preview_rows} rows and {preview_cols} columns)"

                    workbook.close()

                except Exception as e:
                    content += f"\nError reading Excel file: {e}"
            else:
                content += "\n(openpyxl not available for detailed preview)"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_text_document_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate preview for text-based documents"""
        try:
            # Read as text file
            try:
                content_text = file_path.read_text(encoding='utf-8', errors='ignore')
            except UnicodeDecodeError:
                content_text = file_path.read_text(encoding='latin-1', errors='ignore')

            lines = content_text.split('\n')

            metadata = {
                'file_size': file_path.stat().st_size,
                'line_count': len(lines),
                'character_count': len(content_text),
                'word_count': len(content_text.split()),
                'encoding': 'utf-8'
            }

            # Determine document type
            extension = file_path.suffix.lower()
            doc_type = {
                '.md': 'Markdown',
                '.rst': 'reStructuredText',
                '.txt': 'Plain Text'
            }.get(extension, 'Text Document')

            content = f"📄 {doc_type}: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"
            content += f"Lines: {metadata['line_count']:,}\n"
            content += f"Words: {metadata['word_count']:,}\n"
            content += f"Characters: {metadata['character_count']:,}\n"

            # Show content preview
            if content_text.strip():
                content += "\n📝 Content Preview:\n"

                # Truncate if too long
                preview_lines = lines[:self.config.max_lines]
                preview_text = '\n'.join(preview_lines)

                if len(preview_text) > self.config.max_text_length:
                    preview_text = preview_text[:self.config.max_text_length] + "..."

                content += preview_text

                if len(lines) > self.config.max_lines:
                    content += f"\n\n... (showing first {self.config.max_lines} of {len(lines)} lines)"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_csv_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate CSV preview"""
        try:
            import csv

            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': 'text/csv'
            }

            content = f"📊 CSV File: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"

            # Read CSV data
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as csvfile:
                # Detect delimiter
                sample = csvfile.read(1024)
                csvfile.seek(0)

                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter

                reader = csv.reader(csvfile, delimiter=delimiter)
                rows = []

                for i, row in enumerate(reader):
                    rows.append(row)
                    if i >= 20:  # Limit to first 20 rows
                        break

            if rows:
                metadata.update({
                    'estimated_rows': len(rows),
                    'columns': len(rows[0]) if rows else 0,
                    'delimiter': delimiter
                })

                content += f"Columns: {metadata['columns']}\n"
                content += f"Delimiter: '{delimiter}'\n"
                content += f"Sample Rows: {len(rows)}\n"

                content += "\n📋 Data Preview:\n"

                # Show header if available
                if rows:
                    header = rows[0]
                    content += f"Headers: {' | '.join(header[:5])}\n"  # First 5 columns
                    content += "-" * 50 + "\n"

                    # Show data rows
                    for i, row in enumerate(rows[1:11]):  # Next 10 rows
                        row_preview = ' | '.join([str(cell)[:15] for cell in row[:5]])  # First 5 columns, truncated
                        content += f"Row {i+1}: {row_preview}\n"

                    if len(rows) > 11:
                        content += f"... (showing first 10 data rows)"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_generic_document_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate generic document preview"""
        try:
            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': mimetypes.guess_type(str(file_path))[0]
            }

            extension = file_path.suffix.upper().lstrip('.')
            content = f"📄 {extension} Document: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"
            content += f"Type: {metadata.get('mime_type', 'Unknown')}\n"
            content += "\n(Detailed preview not available for this document type)"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.DOCUMENT,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_archive_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate archive preview"""
        try:
            extension = file_path.suffix.lower()

            if extension == '.zip':
                return await self._generate_zip_preview(file_path, start_time)
            elif extension in ['.tar', '.gz', '.tgz', '.bz2', '.tbz2', '.xz']:
                return await self._generate_tar_preview(file_path, start_time)
            else:
                return await self._generate_generic_archive_preview(file_path, start_time)

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_zip_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate ZIP archive preview"""
        try:
            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': 'application/zip'
            }

            content = f"📦 ZIP Archive: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"

            with zipfile.ZipFile(file_path, 'r') as zip_file:
                file_list = zip_file.filelist

                metadata.update({
                    'file_count': len(file_list),
                    'compressed_size': sum(info.compress_size for info in file_list),
                    'uncompressed_size': sum(info.file_size for info in file_list)
                })

                content += f"Files: {metadata['file_count']}\n"
                content += f"Compressed Size: {self._format_file_size(metadata['compressed_size'])}\n"
                content += f"Uncompressed Size: {self._format_file_size(metadata['uncompressed_size'])}\n"

                if metadata['uncompressed_size'] > 0:
                    compression_ratio = (1 - metadata['compressed_size'] / metadata['uncompressed_size']) * 100
                    content += f"Compression Ratio: {compression_ratio:.1f}%\n"

                # List contents
                content += "\n📁 Contents:\n"

                # Group by directory
                directories = set()
                files = []

                for info in file_list[:self.config.archive_max_entries]:
                    if info.is_dir():
                        directories.add(info.filename)
                    else:
                        files.append(info)
                        # Add parent directories
                        parent = str(Path(info.filename).parent)
                        if parent != '.':
                            directories.add(parent + '/')

                # Show directories first
                for directory in sorted(directories):
                    content += f"📁 {directory}\n"

                # Show files
                for info in sorted(files, key=lambda x: x.filename)[:20]:  # First 20 files
                    size_str = self._format_file_size(info.file_size)
                    content += f"📄 {info.filename} ({size_str})\n"

                if len(file_list) > self.config.archive_max_entries:
                    content += f"... and {len(file_list) - self.config.archive_max_entries} more files"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.ARCHIVE,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_tar_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate TAR archive preview"""
        try:
            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': 'application/x-tar'
            }

            extension = file_path.suffix.lower()
            archive_type = {
                '.tar': 'TAR',
                '.gz': 'GZIP',
                '.tgz': 'TAR.GZ',
                '.bz2': 'BZIP2',
                '.tbz2': 'TAR.BZ2',
                '.xz': 'XZ'
            }.get(extension, 'TAR')

            content = f"📦 {archive_type} Archive: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"

            # Determine open mode
            if extension in ['.gz', '.tgz']:
                mode = 'r:gz'
            elif extension in ['.bz2', '.tbz2']:
                mode = 'r:bz2'
            elif extension == '.xz':
                mode = 'r:xz'
            else:
                mode = 'r'

            with tarfile.open(file_path, mode) as tar_file:
                members = tar_file.getmembers()

                total_size = sum(member.size for member in members if member.isfile())

                metadata.update({
                    'file_count': len([m for m in members if m.isfile()]),
                    'directory_count': len([m for m in members if m.isdir()]),
                    'total_size': total_size
                })

                content += f"Files: {metadata['file_count']}\n"
                content += f"Directories: {metadata['directory_count']}\n"
                content += f"Total Size: {self._format_file_size(metadata['total_size'])}\n"

                # List contents
                content += "\n📁 Contents:\n"

                for member in members[:self.config.archive_max_entries]:
                    if member.isdir():
                        content += f"📁 {member.name}/\n"
                    elif member.isfile():
                        size_str = self._format_file_size(member.size)
                        content += f"📄 {member.name} ({size_str})\n"
                    elif member.islnk() or member.issym():
                        content += f"🔗 {member.name} -> {member.linkname}\n"

                if len(members) > self.config.archive_max_entries:
                    content += f"... and {len(members) - self.config.archive_max_entries} more entries"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.ARCHIVE,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    async def _generate_generic_archive_preview(self, file_path: Path, start_time: float) -> PreviewResult:
        """Generate generic archive preview"""
        try:
            metadata = {
                'file_size': file_path.stat().st_size,
                'mime_type': mimetypes.guess_type(str(file_path))[0]
            }

            extension = file_path.suffix.upper().lstrip('.')
            content = f"📦 {extension} Archive: {file_path.name}\n"
            content += f"Size: {self._format_file_size(metadata['file_size'])}\n"
            content += f"Type: {metadata.get('mime_type', 'Unknown')}\n"
            content += "\n(Detailed preview not available for this archive type)"

            return PreviewResult(
                file_path=file_path,
                preview_type=PreviewType.ARCHIVE,
                content=content,
                metadata=metadata,
                generation_time=time.time() - start_time
            )

        except Exception as e:
            return self._create_error_result(file_path, str(e))

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"